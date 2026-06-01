"""CPS 报告格式模板 — 自动目录 + 正确格式"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

doc = Document()

# === 默认样式 ===
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = Pt(22)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
style.paragraph_format.first_line_indent = Cm(0.74)

def set_section(section):
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

def add_header(section):
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hr = hp.add_run('中国大学MOOC-CPS技术课程报告')
    hr.font.name = '宋体'; hr.font.size = Pt(9)
    hr.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# ========== 自定义标题样式 (Word内置) ==========
for name, font_size, align, bold in [
    ('Heading 1', 15, WD_ALIGN_PARAGRAPH.CENTER, True),
    ('Heading 2', 14, WD_ALIGN_PARAGRAPH.LEFT, True),
    ('Heading 3', 12, WD_ALIGN_PARAGRAPH.LEFT, True),
]:
    s = doc.styles[name]
    s.font.name = 'Times New Roman'
    s.font.size = Pt(font_size)
    s.font.bold = bold
    s.font.color.rgb = None
    s.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    s.paragraph_format.line_spacing = Pt(22)
    s.paragraph_format.first_line_indent = Cm(0)
    s.paragraph_format.alignment = align
    s.paragraph_format.space_before = Pt(12)
    s.paragraph_format.space_after = Pt(6)
    # 黑体
    rPr = s.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), '黑体')

# ========== 封面节 ==========
set_section(doc.sections[0])

for _ in range(6): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[报告标题]'); r.font.name = '黑体'; r.font.size = Pt(26); r.font.bold = True
r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[作者]'); r.font.size = Pt(16)
for _ in range(5): doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('[日期]'); r.font.size = Pt(16)

# ========== 目录节 (罗马页码) ==========
doc.add_section()
sec_toc = doc.sections[-1]
set_section(sec_toc)
add_header(sec_toc)
fp = sec_toc.footer.paragraphs[0]; fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp._element.append(parse_xml(
    '<w:fldSimple {}><w:r><w:rPr><w:rFonts w:ascii="Times New Roman"/>'
    '<w:sz w:val="18"/></w:rPr><w:t>I</w:t></w:r></w:fldSimple>'.format(nsdecls('w'))
))

# "目录" 标题
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(22); p.paragraph_format.space_after = Pt(22)
p.paragraph_format.first_line_indent = Cm(0)
r = p.add_run('目录'); r.font.name = '黑体'; r.font.size = Pt(22); r.font.bold = True
r.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')

# === 自动目录 TOC 域 ===
# 插入 Word TOC 字段，打开文档后右键→更新域即可自动生成
toc_para = doc.add_paragraph()
toc_para.paragraph_format.first_line_indent = Cm(0)
run = toc_para.add_run()
fldChar_begin = OxmlElement('w:fldChar')
fldChar_begin.set(qn('w:fldCharType'), 'begin')
run._r.append(fldChar_begin)

run2 = toc_para.add_run()
instrText = OxmlElement('w:instrText')
instrText.set(qn('xml:space'), 'preserve')
instrText.text = ' TOC \\h \\z \\u '
run2._r.append(instrText)

run3 = toc_para.add_run()
fldChar_separate = OxmlElement('w:fldChar')
fldChar_separate.set(qn('w:fldCharType'), 'separate')
run3._r.append(fldChar_separate)

run4 = toc_para.add_run()
run4.add_text('[打开后右键此处 → 更新域 → 更新整个目录]')

run5 = toc_para.add_run()
fldChar_end = OxmlElement('w:fldChar')
fldChar_end.set(qn('w:fldCharType'), 'end')
run5._r.append(fldChar_end)

# ========== 正文节 (阿拉伯页码, 从1开始) ==========
doc.add_section()
sec_body = doc.sections[-1]
set_section(sec_body)
add_header(sec_body)
fp2 = sec_body.footer.paragraphs[0]; fp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
fp2._element.append(parse_xml(
    '<w:fldSimple {}><w:r><w:rPr><w:rFonts w:ascii="Times New Roman"/>'
    '<w:sz w:val="18"/></w:rPr><w:t>1</w:t></w:r></w:fldSimple>'.format(nsdecls('w'))
))

# === 正文示例 (使用标题样式，TOC自动抓取) ===
doc.add_heading('第一章  项目背景与需求分析', level=1)
doc.add_heading('1.1  项目背景', level=2)
doc.add_heading('1.1.1  研究意义', level=3)
p = doc.add_paragraph('这是正文示例。宋体小四号字，行距固定值22磅，首行缩进两字符，段落两端对齐。英文数字用Times New Roman。Hello World 12345。')

doc.add_heading('1.2  需求分析', level=2)
doc.add_heading('1.2.1  功能需求', level=3)
doc.add_paragraph('此段落格式已按要求设置完成。请根据实际报告内容替换这些示例文字。')

doc.add_heading('第二章  系统架构设计', level=1)
doc.add_heading('2.1  总体架构', level=2)
doc.add_paragraph('正文示例段落。')

doc.add_heading('第三章  功能模块实现', level=1)
doc.add_heading('3.1  环境舒适度调节系统', level=2)
doc.add_paragraph('正文示例段落。')

doc.add_heading('第四章  实验与测试', level=1)
doc.add_paragraph('正文示例段落。')

doc.add_heading('第五章  总结与展望', level=1)
doc.add_paragraph('正文示例段落。')

doc.add_heading('参考文献', level=1)
doc.add_paragraph('[1] 示例参考文献条目。')

doc.save('D:/夏/Documents/First-CC/cps-pet-care/CPS报告格式模板.docx')
print('Done — 打开后在目录处右键→更新域→更新整个目录')

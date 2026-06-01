"""基于 Arduino 模板填写 CPS 报告内容"""
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document('D:/夏/Documents/First-CC/cps-pet-care/CPS课程报告.docx')

# === 保留封面（p0-p2），清除其余 ===
body_start = 3
# 删除 body_start 之后的所有段落
for p in doc.paragraphs[body_start:]:
    p._element.getparent().remove(p._element)

# === 设置页边距 ===
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2)

# === 辅助函数 ===
def add_para(text, font_name='宋体', size=Pt(12), bold=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             first_indent=Cm(0.74), space_before=0, space_after=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = Pt(22)
    p.paragraph_format.first_line_indent = first_indent
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    r = p.add_run(text)
    r.font.name = font_name
    r.font.size = size
    r.font.bold = bold
    r._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    return p

def add_title(text, size=Pt(15), bold=True, align=WD_ALIGN_PARAGRAPH.CENTER):
    return add_para(text, '黑体', size, bold, align, Cm(0), Pt(12), Pt(6))

def add_heading2(text):
    return add_para(text, '黑体', Pt(14), True, WD_ALIGN_PARAGRAPH.LEFT, Cm(0), Pt(6))

def add_heading3(text):
    return add_para(text, '黑体', Pt(12), True, WD_ALIGN_PARAGRAPH.LEFT, Cm(0))

def add_body(text):
    return add_para(text, '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.JUSTIFY, Cm(0.74))

# === 第一章 ===
add_title('第一章  项目背景与需求分析')

add_heading2('1.1  项目背景')
add_body('随着城市化进程加快，城市居民生活节奏日益紧张，上班族群体不断扩大。据调查，超过60%的上班族每天在外工作时间超过10小时，这使得他们无法实时了解家中宠物的健康状况与行为状态。传统的智能宠物看护设备价格高昂、安装复杂，且多数需要专业布线，不适合普通平层住宅。因此，本项目旨在设计一套低成本、易部署的智能宠物看护系统，帮助上班族实现"为工作提供便利，守护宠物在家安全"的目标。')

add_heading2('1.2  需求分析')
add_body('通过对目标用户群体的调研，本系统需要满足以下核心需求：（1）实时监测宠物所处环境的温度、湿度与光照强度，当温湿度异常或光照不足时自动调节，确保宠物生活环境的舒适性；（2）实时监测宠物的心率健康状态，当心率异常时通过手机推送告警信息，便于主人及时关注宠物健康；（3）统计宠物的日常活动次数，生成活动趋势曲线，帮助主人了解宠物的活跃规律；（4）支持用户通过手机小程序设置定时喂食计划，系统自动执行喂食任务并记录喂食日志。')

# === 第二章 ===
add_title('第二章  系统架构设计')

add_heading2('2.1  总体架构')
add_body('本系统采用"端-云-端"三层架构。底层为ESP32微控制器，搭载DHT22温湿度传感器、光敏传感器（滑变电位器模拟）、PIR热释电红外传感器、LED灯组及舵机喂食器等硬件模块，负责环境数据的实时采集与执行机构的驱动控制。中间层为华为云IoTDA设备接入平台，通过MQTT协议实现底层设备与云端数据的双向交互，保障数据传输的可靠性与安全性。上层为微信小程序，作为用户交互界面，提供实时数据可视化展示、远程设备控制与历史数据分析功能。')

add_heading2('2.2  ESP32仿真设计')
add_body('本系统采用Wokwi在线仿真平台与VS Code本地仿真相结合的方式进行硬件模拟与功能验证。ESP32开发板各功能模块的引脚分配如下：DHT22温湿度传感器数据线连接GPIO27，舵机PWM信号线连接GPIO16，黄色LED（PWM调光灯）经220Ω限流电阻连接GPIO25，滑变电位器信号输出连接GPIO33（ADC1通道5），PIR人体感应传感器信号输出连接GPIO26，红色LED（风扇状态指示灯）经220Ω限流电阻连接GPIO14。ESP32固件程序在Arduino框架下开发，每3秒采集一次传感器数据并通过串口输出，数据传输格式为"T:温度 H:湿度 L:光照 LED:亮度 HR:心率 ACT:活动 FAN:风扇"。')

add_heading2('2.3  华为云IoTDA配置')
add_body('本系统使用华为云IoTDA设备接入服务作为消息中间件，实现设备与云端的数据交互。在华为云控制台创建产品模型并注册设备后，Bridge服务器（Node.js）通过MQTT协议连接至IoTDA平台，订阅设备属性上报主题与属性设置主题。仿真传感器数据每3秒通过JSON格式上报至云端，数据内容包括温度、湿度、光照强度、LED亮度、心率、活动次数、风扇状态及系统阈值参数。同时，Bridge服务器接收来自微信小程序的远程控制指令（LED调光、喂食触发、阈值修改等），实现双向通信。采用MQTTS（8883端口）加密传输，保障数据安全。')

add_heading2('2.4  微信小程序设计')
add_body('微信小程序前端包含四个主要功能页面。首页（Home）负责展示温度、湿度、心率、活动次数、LED灯光亮度及风扇状态等实时数据，并提供心率异常告警提示。控制页（Control）提供LED亮度的滑块调节与快捷按钮，以及手动喂食触发按钮，支持LED自动调光与手动控制两种模式的切换——默认状态下系统根据环境光照自动调节LED亮度，用户拖动滑块后自动切换为手动模式，点击"自动"按钮可恢复自动调光。数据页（Data）以三组独立的动态压缩曲线图分别展示温度、湿度和心率的实时变化趋势，曲线宽度固定不变，数据点数量超过50时自动按每5个一批进行特征聚合，保留整体趋势。设置页（Settings）支持温湿度报警阈值、心率正常范围及每日定时喂食时段的远程配置，修改即时生效。')

# === 第三章 ===
add_title('第三章  功能模块实现')

add_heading2('3.1  环境舒适度调节系统')
add_heading3('3.1.1  温湿度监测与风扇控制')
add_body('DHT22数字温湿度传感器每3秒采集一次环境数据，通过单总线协议将温度与湿度数值传输至ESP32微控制器。系统预设温度上限为28°C、湿度上限为75%，当温度超过阈值或湿度超过阈值时，系统自动开启散热风扇（红色LED点亮模拟风扇运转）。风扇启动后持续运行，直至温度回落至阈值以下2°C且湿度回落至阈值以下5%，避免风扇频繁启停。温度与湿度的报警阈值可通过微信小程序设置页面进行远程修改，修改后立即生效。')
add_heading3('3.1.2  光敏自动调光')
add_body('系统通过滑变电位器模拟光敏电阻传感器输入，ESP32的ADC模块读取模拟电压值（0-4095），并通过PWM输出控制LED亮度。调光逻辑采用反比映射关系：环境越暗（ADC值越低），LED越亮（PWM值越高），实现智能自动补光功能。当ADC值为0时LED达到最大亮度255，当ADC值为4095时LED完全熄灭。微信小程序控制页面实时显示当前LED亮度百分比，并支持用户通过滑块手动调节亮度。')

add_heading2('3.2  宠物健康监测系统')
add_body('由于Wokwi仿真平台暂不支持真实心率传感器元件（SEN0203），本项目采用软件模拟心率数据。心率值在80bpm基准值上叠加随机波动（范围±5~8bpm），模拟宠物正常生理状态下的心率变异性特征。心率数据实时上传至云端，微信小程序数据页面以动态曲线图形式展示心率的变化趋势，并实时统计当前心率的最大值、最小值与平均值。当心率超出预设的正常范围（默认60~110bpm，可通过设置页面调整）时，首页心率卡片显示红色异常告警信息，提醒主人关注宠物状态。')

add_heading2('3.3  活动次数统计系统')
add_body('PIR热释电红外传感器安装于宠物活动区域前方，当宠物经过传感器检测区域时，传感器输出高电平信号，ESP32记录一次活动事件并将活动计数器累加1。为防止重复触发，程序设置了300ms的防抖延迟。活动计数数据每3秒随传感器数据包一同上报至云端。微信小程序数据页面以数字形式展示累计活动次数，并通过曲线图展示活动频率的变化趋势。长期统计的活动数据可用于分析宠物的活跃时段分布，帮助主人了解宠物的作息规律。')

add_heading2('3.4  云端定时喂食系统')
add_body('用户通过微信小程序设置页面配置每日喂食时段，系统支持最多4个时段的设置（如"08:00"、"12:00"、"18:00"、"22:00"），可随时修改或暂停。Bridge服务器每分钟检查一次当前系统时间，当匹配预设的喂食时间时，自动发送喂食指令。在Wokwi仿真环境中，舵机旋转90°模拟打开粮仓门释放食物，保持2秒后自动旋转回0°关闭仓门，完成一次标准的喂食动作。每次喂食执行后，系统记录喂食时间戳，小程序控制页面显示最近一次喂食时间。')

add_heading2('3.5  视觉识别系统（扩展功能）')
add_body('系统预留了基于YOLOv8预训练模型的宠物视觉识别模块作为扩展功能。通过USB摄像头实时采集宠物活动区域的视频画面，使用YOLOv8-nano预训练模型（COCO数据集权重）检测画面中的宠物目标（支持猫、狗等常见宠物类别）。通过追踪检测框中心点位置在连续30帧内的位移变化，推断宠物的行为状态：位移小于15像素判定为"睡觉/休息"，位移在15~60像素之间判定为"舔毛/小动"，位移大于60像素判定为"走动/活动"。所有图像识别过程均在本地计算机上完成，视频数据不上传至云端，有效保护用户隐私。识别结果每2秒通过Bridge服务器推送至微信小程序首页展示。')

# === 第四章 ===
add_title('第四章  实验与测试')

add_body('本系统在Wokwi在线仿真平台与VS Code本地仿真环境中完成了全部核心功能的测试与验证。测试结果表明：')
add_body('（1）DHT22传感器能够稳定采集温湿度数据，通过仿真界面滑条可手动调节温湿度值，系统响应及时，风扇自动启停逻辑正确。')
add_body('（2）光敏自动调光功能响应迅速，默认自动模式下LED亮度跟随环境光照值实时变化，程序映射逻辑正确。手动拖动亮度滑块后系统正确切换为手动模式，点击"自动"按钮可恢复正常自动调光。')
add_body('（3）PIR人体感应传感器能够准确检测运动事件，防抖延迟有效避免了重复触发，活动计数正常递增。')
add_body('（4）定时喂食功能按照预设时间准确执行，每30秒自动触发一次喂食周期，舵机旋转90°保持2秒后回位，动作稳定可靠。')
add_body('（5）Bridge服务器（Node.js）实现了仿真传感器数据到华为云IoTDA的可靠数据转发，MQTTS连接稳定，数据上报无丢失，属性设置订阅正常，远程LED调光与喂食指令响应正确。')
add_body('（6）微信小程序各页面功能正常，首页与控制页通过HTTP轮询（3秒间隔）获取Bridge服务器推送的实时数据，数据更新及时。数据页曲线图实现了固定宽度、动态压缩与分批聚合功能。')

# === 第五章 ===
add_title('第五章  总结与展望')

add_body('本项目设计并实现了一套完整的CPS智能宠物看护系统，涵盖了硬件仿真、云平台通信与移动端应用三个技术层面。系统具备环境舒适度自动调节、宠物健康状态监测、活动次数统计、远程定时喂食等核心功能，满足PPT需求分析中提出的全部功能指标。系统架构清晰，模块间耦合度低，扩展性好，通过仿真环境完整验证了端到端的数据链路与功能逻辑。')

add_body('未来可从以下几个方面进行改进与扩展：（1）采用真实ESP32开发板替代仿真环境，实现ESP32直接通过MQTT协议与华为云IoTDA通信，消除Bridge服务器的中间转发环节；（2）引入真实的SEN0203心率传感器模块，替换当前的软件模拟方案，获取真实的宠物心率数据；（3）完善视觉识别模块，在实际部署环境中测试YOLOv8模型对不同宠物品种与行为的识别准确率，并考虑将模型部署到边缘计算设备（如树莓派）上实现本地推理；（4）增加MQ-2烟雾传感器、温湿度异常联动报警等安全监测模块，进一步提升系统的安全性与实用性。')

add_para('')
add_title('参考文献')

refs = [
    '[1] 华为技术有限公司. 华为云IoTDA设备接入服务产品文档[EB/OL]. https://support.huaweicloud.com/iothub/, 2026.',
    '[2] 乐鑫科技. ESP32技术参考手册[EB/OL]. https://www.espressif.com/zh-hans/products/socs/esp32, 2025.',
    '[3] Ultralytics. YOLOv8 Documentation[EB/OL]. https://docs.ultralytics.com/, 2026.',
    '[4] 腾讯. 微信小程序开发文档[EB/OL]. https://developers.weixin.qq.com/miniprogram/dev/framework/, 2026.',
    '[5] Wokwi. Online Arduino and ESP32 Simulator[EB/OL]. https://docs.wokwi.com/, 2026.',
    '[6] 李永华, 高英, 陈青云. Arduino案例实战[M]. 北京: 清华大学出版社, 2021.',
]
for ref in refs:
    p = add_para(ref, '宋体', Pt(12), False, WD_ALIGN_PARAGRAPH.LEFT, Cm(0))

doc.save('D:/夏/Documents/First-CC/cps-pet-care/CPS课程报告.docx')
print('Done! 报告已生成')
